# -*- coding: utf-8 -*-
"""
@Project : copilot_demo
@File    : files_parse.py
@Date    : 2025/12/17 11:12
@Desc    : 
"""
import os

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import pdfplumber


def get_heading_level(paragraph):
    """判断段落是否为标题并返回级别"""
    try:
        style_name = paragraph.style.original_name
    except AttributeError:
        return None
    if style_name.startswith('Heading'):
        try:
            return int(style_name.split()[1])
        except (IndexError, ValueError):
            return None
    elif '标题' in style_name:
        try:
            for char in style_name:
                if char.isdigit():
                    return int(char)
        except ValueError:
            return None
    return None


def get_heading_numbering(paragraph):
    """获取段落的自动编号（如1.1、1.2），兼容不同版本的python-docx"""
    if not hasattr(paragraph._element, 'pPr') or not paragraph._element.pPr:
        return ""

    num_pr = paragraph._element.pPr.numPr
    if not num_pr:
        return ""

    try:
        # 获取编号ID
        num_id = num_pr.numId.val
        # 获取编号定义（兼容不同版本的API）
        numbering_part = paragraph.part.numbering_part
        if not numbering_part:
            return ""

        num = numbering_part.numbering_definitions.get_num(num_id)
        if not num:
            return ""

        # 获取当前级别
        level = num_pr.ilvl.val
        # 获取编号文本
        numbering_text = num.format_for_level(level)
        return numbering_text + " " if numbering_text else ""
    except Exception as e:
        # 发生错误时返回空，避免程序中断
        return ""


class FileParser:
    @classmethod
    def parse_files(cls, file_path):
        ext_with_dot = str(os.path.splitext(file_path)[1]).lower()
        if ext_with_dot == ".docx":
            file_content = cls.parse_docx(file_path)
        elif ext_with_dot == ".pdf":
            file_content = cls.parse_pdf(file_path)
        elif ext_with_dot in [".txt", ".md"]:
            file_content = cls.parse_txt(file_path)
        else:
            file_content = ""

        return file_content

    @staticmethod
    def parse_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        return text_content

    @staticmethod
    def parse_pdf(file_path):
        text_content = ""
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if not page.images:
                    text_content += page.extract_text() + "\n"
        return text_content

    @staticmethod
    def parse_docx(file_path):

        """提取并格式化Word文档内容，包含自动生成的标题编号"""
        doc = docx.Document(file_path)
        ordered_content = []

        for element in doc.element.body:
            # 处理段落（包括标题）
            if isinstance(element, CT_P):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    text = para.text
                    if text:
                        heading_level = get_heading_level(para)
                        # 获取自动编号
                        numbering = get_heading_numbering(para) if heading_level else ""

                        if heading_level:
                            # 标题格式：编号 + 内容
                            ordered_content.append(f"\n{'#' * heading_level} {numbering}{text}\n")
                        else:
                            # 普通文本
                            ordered_content.append(f"{text}\n")

            # 处理表格
            elif isinstance(element, CT_Tbl):
                table_index = None
                found_table = None
                for idx, tbl in enumerate(doc.tables):
                    if tbl._element == element:
                        table_index = idx
                        found_table = tbl
                        break

                if found_table and table_index is not None:
                    # 提取表格内容
                    for row in found_table.rows:
                        row_data = [cell.text for cell in row.cells]
                        ordered_content.append("| ".join(row_data) + " |\n")

        return ''.join(ordered_content)
